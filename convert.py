"""
Convert Ocean_From_Motion.docx into a static HTML book in docs/.
Handles Word OMML math → LaTeX rendered via KaTeX.
"""

import os
import re
import shutil
from html import escape
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOCX_PATH = os.path.join(os.path.dirname(__file__),
                         "Ocean_From_Motion (1).docx")
DOCS_DIR = os.path.join(os.path.dirname(__file__), "docs")
BASE_URL = "/Ocean-From-Motion"

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── OMML to LaTeX ───────────────────────────────────────────────────

def _mtag(local):
    return f"{{{MATH_NS}}}{local}"

def _wtag(local):
    return f"{{{WORD_NS}}}{local}"


# Map common Unicode math chars to LaTeX commands
UNICODE_TO_LATEX = {
    "κ": r"\kappa",
    "φ": r"\varphi",
    "δ": r"\delta",
    "Σ": r"\Sigma",
    "σ": r"\sigma",
    "π": r"\pi",
    "α": r"\alpha",
    "β": r"\beta",
    "γ": r"\gamma",
    "ε": r"\varepsilon",
    "ζ": r"\zeta",
    "η": r"\eta",
    "θ": r"\theta",
    "λ": r"\lambda",
    "μ": r"\mu",
    "ν": r"\nu",
    "ξ": r"\xi",
    "ρ": r"\rho",
    "τ": r"\tau",
    "ω": r"\omega",
    "Ω": r"\Omega",
    "Δ": r"\Delta",
    "Φ": r"\Phi",
    "Γ": r"\Gamma",
    "Λ": r"\Lambda",
    "Π": r"\Pi",
    "Θ": r"\Theta",
    "∈": r"\in",
    "∉": r"\notin",
    "⊂": r"\subset",
    "⊆": r"\subseteq",
    "⊃": r"\supset",
    "∪": r"\cup",
    "∩": r"\cap",
    "∅": r"\emptyset",
    "∞": r"\infty",
    "≤": r"\leq",
    "≥": r"\geq",
    "≠": r"\neq",
    "≈": r"\approx",
    "≡": r"\equiv",
    "±": r"\pm",
    "×": r"\times",
    "÷": r"\div",
    "·": r"\cdot",
    "→": r"\to",
    "←": r"\leftarrow",
    "↔": r"\leftrightarrow",
    "⇒": r"\Rightarrow",
    "⇐": r"\Leftarrow",
    "⇔": r"\Leftrightarrow",
    "∀": r"\forall",
    "∃": r"\exists",
    "¬": r"\neg",
    "∧": r"\wedge",
    "∨": r"\vee",
    "⊕": r"\oplus",
    "⊗": r"\otimes",
    "ℕ": r"\mathbb{N}",
    "ℤ": r"\mathbb{Z}",
    "ℝ": r"\mathbb{R}",
    "ℚ": r"\mathbb{Q}",
    "ℂ": r"\mathbb{C}",
    "′": r"'",
    "″": r"''",
    "∣": r"|",
    "⌀": r"\emptyset",
    "∼": r"\sim",
    "≾": r"\precsim",
    "𝜅": r"\kappa",
    "\u2061": "",  # function application (invisible)
    "\u200b": "",  # zero-width space
    "\u200c": "",  # zero-width non-joiner
    "\u2009": r"\,",  # thin space
    "\u2005": r"\;",  # four-per-em space
}


def _latex_escape_text(text):
    """Escape a text string for LaTeX, converting Unicode math symbols.
    Ensures spacing between LaTeX commands and following letters."""
    out = []
    for ch in text:
        if ch in UNICODE_TO_LATEX:
            replacement = UNICODE_TO_LATEX[ch]
            # If previous output ends with \command and this replacement starts
            # with a letter (or is a letter), add space
            if out and replacement:
                prev = out[-1]
                if re.search(r'\\[a-zA-Z]+$', prev) and replacement[0].isalpha():
                    out.append(" ")
            out.append(replacement)
        elif ch in "#$%&_{}":
            out.append("\\" + ch)
        elif ch == "~":
            out.append(r"\sim")
        elif ch == "^":
            out.append(r"\hat{}")
        elif ch == "\\":
            out.append(r"\backslash")
        else:
            # If previous output ends with \command and this char is a letter
            if out and ch.isalpha():
                prev = out[-1]
                if re.search(r'\\[a-zA-Z]+$', prev):
                    out.append(" ")
            out.append(ch)
    return "".join(out)


def _join_latex_parts(parts):
    """Join LaTeX parts, adding spaces where a command would run into a letter."""
    result = []
    for i, part in enumerate(parts):
        if not part:
            continue
        # If previous part ends with a backslash-command (letters) and this part
        # starts with a letter, add a space to prevent them from merging.
        if result:
            prev = result[-1]
            # Check if prev ends with \command (letters)
            if re.search(r'\\[a-zA-Z]+$', prev) and part and part[0].isalpha():
                result.append(" ")
            # Also if prev ends with a letter and current starts with a letter
            # but they are separate OMML elements (subscripts, etc.), add thin space
        result.append(part)
    return "".join(result)


def omml_to_latex(elem):
    """Recursively convert an OMML element to a LaTeX string."""
    tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

    if tag == "oMathPara":
        # Display math paragraph — convert inner oMath
        parts = []
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "oMath":
                parts.append(omml_to_latex(child))
        return "".join(parts)

    if tag == "oMath":
        parts = [omml_to_latex(c) for c in elem]
        return _join_latex_parts(parts)

    if tag == "r":
        # Math run — extract text from m:t
        text = ""
        is_normal = False  # roman/normal style
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "t":
                text += child.text or ""
            elif ct == "rPr":
                for prop in child:
                    pt = prop.tag.split("}")[-1] if "}" in prop.tag else prop.tag
                    if pt == "sty":
                        val = prop.get(f"{{{MATH_NS}}}val", "")
                        if val == "p":  # plain/roman
                            is_normal = True
        latex = _latex_escape_text(text)
        # For plain/roman style: only wrap sequences of 2+ letters in \text{}
        # to render them upright. Don't wrap operators, LaTeX commands, or symbols.
        if is_normal and latex.strip():
            # Check if it's purely alphabetic multi-letter text (like "where", "for")
            stripped = text.strip()
            if len(stripped) > 1 and stripped.isalpha():
                latex = r"\text{" + stripped + "}"
        return latex

    if tag == "sSub":
        # Subscript: e_{sub}
        base = ""
        sub = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = omml_to_latex(child)
            elif ct == "sub":
                sub = omml_to_latex(child)
            # skip sSubPr
        return f"{base}_{{{sub}}}"

    if tag == "sSup":
        # Superscript: e^{sup}
        base = ""
        sup = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = omml_to_latex(child)
            elif ct == "sup":
                sup = omml_to_latex(child)
        return f"{base}^{{{sup}}}"

    if tag == "sSubSup":
        # Sub+Superscript: e_{sub}^{sup}
        base = ""
        sub = ""
        sup = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                base = omml_to_latex(child)
            elif ct == "sub":
                sub = omml_to_latex(child)
            elif ct == "sup":
                sup = omml_to_latex(child)
        return f"{base}_{{{sub}}}^{{{sup}}}"

    if tag == "d":
        # Delimited group (parentheses, brackets, abs value, etc.)
        beg_chr = "("
        end_chr = ")"
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "dPr":
                for prop in child:
                    pt = prop.tag.split("}")[-1] if "}" in prop.tag else prop.tag
                    if pt == "begChr":
                        beg_chr = prop.get(f"{{{MATH_NS}}}val", "(")
                    elif pt == "endChr":
                        end_chr = prop.get(f"{{{MATH_NS}}}val", ")")

        # Map delimiter characters
        delim_map = {
            "(": r"\left(", ")": r"\right)",
            "[": r"\left[", "]": r"\right]",
            "{": r"\left\{", "}": r"\right\}",
            "|": r"\left|", "⟨": r"\left\langle", "⟩": r"\right\rangle",
            "‖": r"\left\|",
        }
        beg_latex = delim_map.get(beg_chr, r"\left" + beg_chr)
        end_latex = delim_map.get(end_chr, r"\right" + end_chr)
        if end_chr == "|":
            end_latex = r"\right|"
        if end_chr == "‖":
            end_latex = r"\right\|"

        # Collect the e (element) children
        parts = []
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                parts.append(omml_to_latex(child))
        inner = ", ".join(parts) if len(parts) > 1 else "".join(parts)
        return f"{beg_latex}{inner}{end_latex}"

    if tag == "f":
        # Fraction
        num = ""
        den = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "num":
                num = omml_to_latex(child)
            elif ct == "den":
                den = omml_to_latex(child)
        return rf"\frac{{{num}}}{{{den}}}"

    if tag == "nary":
        # N-ary operator (sum, product, etc.)
        char = "∑"
        sub_val = ""
        sup_val = ""
        body = ""
        sup_hide = False
        sub_hide = False
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "naryPr":
                for prop in child:
                    pt = prop.tag.split("}")[-1] if "}" in prop.tag else prop.tag
                    if pt == "chr":
                        char = prop.get(f"{{{MATH_NS}}}val", "∑")
                    elif pt == "supHide":
                        sup_hide = prop.get(f"{{{MATH_NS}}}val", "0") == "1"
                    elif pt == "subHide":
                        sub_hide = prop.get(f"{{{MATH_NS}}}val", "0") == "1"
            elif ct == "sub":
                sub_val = omml_to_latex(child)
            elif ct == "sup":
                sup_val = omml_to_latex(child)
            elif ct == "e":
                body = omml_to_latex(child)

        nary_map = {"∑": r"\sum", "∏": r"\prod", "∫": r"\int",
                     "∮": r"\oint", "⋃": r"\bigcup", "⋂": r"\bigcap"}
        op = nary_map.get(char, r"\sum")
        result = op
        if sub_val and not sub_hide:
            result += f"_{{{sub_val}}}"
        if sup_val and not sup_hide:
            result += f"^{{{sup_val}}}"
        result += f" {body}"
        return result

    if tag == "acc":
        # Accent (hat, bar, tilde, etc.)
        char = "\u0302"  # combining circumflex (hat) default
        body = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "accPr":
                for prop in child:
                    pt = prop.tag.split("}")[-1] if "}" in prop.tag else prop.tag
                    if pt == "chr":
                        char = prop.get(f"{{{MATH_NS}}}val", "\u0302")
            elif ct == "e":
                body = omml_to_latex(child)

        acc_map = {
            "\u0302": "hat", "\u0300": "grave", "\u0301": "acute",
            "\u0303": "tilde", "\u0304": "bar", "\u0307": "dot",
            "\u0308": "ddot", "\u20d7": "vec", "̂": "hat",
            "̃": "tilde", "̄": "bar", "→": "vec",
        }
        cmd = acc_map.get(char, "hat")
        return rf"\{cmd}{{{body}}}"

    if tag == "rad":
        # Radical (square root, nth root)
        deg = ""
        body = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "deg":
                deg = omml_to_latex(child)
            elif ct == "e":
                body = omml_to_latex(child)
        if deg.strip():
            return rf"\sqrt[{deg}]{{{body}}}"
        return rf"\sqrt{{{body}}}"

    if tag == "limUpp":
        # Upper limit
        body = ""
        lim = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                body = omml_to_latex(child)
            elif ct == "lim":
                lim = omml_to_latex(child)
        return rf"\overset{{{lim}}}{{{body}}}"

    if tag == "limLow":
        # Lower limit
        body = ""
        lim = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                body = omml_to_latex(child)
            elif ct == "lim":
                lim = omml_to_latex(child)
        return rf"\underset{{{lim}}}{{{body}}}"

    if tag == "func":
        # Function application (sin, cos, lim, etc.)
        fname = ""
        body = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "fName":
                fname = omml_to_latex(child).strip()
            elif ct == "e":
                body = omml_to_latex(child)
        known_funcs = {"sin", "cos", "tan", "log", "ln", "exp", "lim",
                       "max", "min", "sup", "inf", "det", "gcd"}
        if fname.replace("\\mathrm{", "").replace("}", "") in known_funcs:
            clean = fname.replace("\\mathrm{", "").replace("}", "")
            return rf"\{clean} {body}"
        return rf"\operatorname{{{fname}}} {body}"

    if tag == "eqArr":
        # Equation array
        lines = []
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                lines.append(omml_to_latex(child))
        return r" \\ ".join(lines)

    if tag == "groupChr":
        # Group character (underbrace, overbrace)
        body = ""
        char = "⏟"
        pos = "bot"
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "groupChrPr":
                for prop in child:
                    pt = prop.tag.split("}")[-1] if "}" in prop.tag else prop.tag
                    if pt == "chr":
                        char = prop.get(f"{{{MATH_NS}}}val", "⏟")
                    elif pt == "pos":
                        pos = prop.get(f"{{{MATH_NS}}}val", "bot")
            elif ct == "e":
                body = omml_to_latex(child)
        if pos == "top" or char == "⏞":
            return rf"\overbrace{{{body}}}"
        return rf"\underbrace{{{body}}}"

    if tag == "m":
        # Matrix
        rows = []
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "mr":
                cols = []
                for mc in child:
                    mct = mc.tag.split("}")[-1] if "}" in mc.tag else mc.tag
                    if mct == "e":
                        cols.append(omml_to_latex(mc))
                rows.append(" & ".join(cols))
        return r"\begin{pmatrix} " + r" \\ ".join(rows) + r" \end{pmatrix}"

    if tag == "box":
        body = ""
        for child in elem:
            ct = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if ct == "e":
                body = omml_to_latex(child)
        return body

    # Container elements: e, num, den, sub, sup, deg, lim, fName
    if tag in ("e", "num", "den", "sub", "sup", "deg", "lim", "fName",
               "oMathParaPr", "ctrlPr", "sSubPr", "sSupPr", "sSubSupPr",
               "dPr", "fPr", "naryPr", "accPr", "radPr", "funcPr",
               "eqArrPr", "groupChrPr", "mPr", "boxPr", "limUppPr", "limLowPr",
               "rPr"):
        if tag in ("e", "num", "den", "sub", "sup", "deg", "lim", "fName"):
            parts = [omml_to_latex(c) for c in elem]
            return _join_latex_parts(parts)
        return ""  # properties — skip

    # Fallback: recurse
    parts = [omml_to_latex(c) for c in elem]
    return _join_latex_parts(parts)


# ── Paragraph to HTML (with math) ──────────────────────────────────

def _para_xml_to_html(para_elem, paragraph):
    """
    Walk the paragraph's XML directly to interleave text runs and math.
    Returns an HTML string with inline LaTeX \\(...\\) or display \\[...\\].
    """
    parts = []
    has_display_math = False

    for child in para_elem:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "r":
            # Regular word run
            text = ""
            bold = False
            italic = False
            for rc in child:
                rt = rc.tag.split("}")[-1] if "}" in rc.tag else rc.tag
                if rt == "t":
                    text += rc.text or ""
                elif rt == "rPr":
                    for rp in rc:
                        rpt = rp.tag.split("}")[-1] if "}" in rp.tag else rp.tag
                        if rpt == "b":
                            val = rp.get(f"{{{WORD_NS}}}val", "true")
                            bold = val != "false" and val != "0"
                        elif rpt == "i":
                            val = rp.get(f"{{{WORD_NS}}}val", "true")
                            italic = val != "false" and val != "0"
            if text:
                t = escape(text)
                if bold and italic:
                    t = f"<strong><em>{t}</em></strong>"
                elif bold:
                    t = f"<strong>{t}</strong>"
                elif italic:
                    t = f"<em>{t}</em>"
                parts.append(t)

        elif tag == "oMathPara":
            # Display math block
            latex = omml_to_latex(child).strip()
            if latex:
                parts.append(f'\\[{latex}\\]')
                has_display_math = True

        elif tag == "oMath":
            # Inline math (oMath directly in paragraph, not inside oMathPara)
            latex = omml_to_latex(child).strip()
            if latex:
                parts.append(f'\\({latex}\\)')

    content = "".join(parts)
    return content, has_display_math


def runs_to_html(paragraph):
    """Convert a paragraph's runs into HTML, preserving bold/italic."""
    parts = []
    for run in paragraph.runs:
        text = escape(run.text)
        if not text:
            continue
        if run.bold and run.italic:
            text = f"<strong><em>{text}</em></strong>"
        elif run.bold:
            text = f"<strong>{text}</strong>"
        elif run.italic:
            text = f"<em>{text}</em>"
        parts.append(text)
    return "".join(parts) or escape(paragraph.text)


def paragraph_to_html(paragraph):
    """Convert a paragraph to HTML, handling math if present."""
    para_elem = paragraph._element
    has_math = any(
        (MATH_NS in c.tag) for c in para_elem
    )

    if has_math:
        content, is_display = _para_xml_to_html(para_elem, paragraph)
        return content, is_display
    else:
        return runs_to_html(paragraph), False


# ── Other helpers ───────────────────────────────────────────────────

def table_to_html(table):
    """Convert a docx table to an HTML <table>."""
    rows_html = []
    for ri, row in enumerate(table.rows):
        cells = []
        tag = "th" if ri == 0 else "td"
        for cell in row.cells:
            cells.append(f"<{tag}>{escape(cell.text)}</{tag}>")
        rows_html.append("<tr>" + "".join(cells) + "</tr>")
    return "<table>\n" + "\n".join(rows_html) + "\n</table>"


def is_section_heading(paragraph):
    """Detect numbered sub-section headings like '4.1 Title' in Body Text."""
    text = paragraph.text.strip()
    if not text:
        return False
    if re.match(r"^\d+(\.\d+)*\.?\s", text):
        if paragraph.runs and all(r.bold for r in paragraph.runs if r.text.strip()):
            return True
        if paragraph.style.name == "Body Text":
            return True
    return False


def classify_heading_level(text):
    """Determine h2/h3/h4 level from numbered heading text."""
    m = re.match(r"^(\d+)(?:\.(\d+))?(?:\.(\d+))?\.", text.strip())
    if not m:
        m = re.match(r"^(\d+)(?:\.(\d+))?(?:\.(\d+))?\s", text.strip())
    if m:
        if m.group(3):
            return "h4"
        elif m.group(2):
            return "h3"
    return "h2"


# ── Document parsing ────────────────────────────────────────────────

def parse_document(doc):
    """Walk the document body, producing ('paragraph', para) or ('table', table)."""
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    elements = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            p = para_map.get(id(child))
            if p:
                elements.append(("paragraph", p))
        elif tag == "tbl":
            t = table_map.get(id(child))
            if t:
                elements.append(("table", t))
    return elements


def split_into_sections(elements):
    """Split elements into preface, chapters, and appendices."""
    sections = {}

    para_idx = 0
    elem_indices = {}
    for ei, (etype, obj) in enumerate(elements):
        if etype == "paragraph":
            elem_indices[para_idx] = ei
            para_idx += 1

    heading1_indices = []
    para_idx = 0
    for ei, (etype, obj) in enumerate(elements):
        if etype == "paragraph":
            if obj.style.name.startswith("Heading"):
                heading1_indices.append((para_idx, ei, obj.text))
            para_idx += 1

    preface_start = elem_indices.get(77, 0)
    preface_end = elem_indices.get(92, 0)
    sections["preface"] = elements[preface_start:preface_end]

    chapter_starts = []
    current_part = None
    for para_idx_val, ei, text in heading1_indices:
        text_stripped = text.strip()
        if text_stripped.startswith("PART"):
            current_part = text_stripped
            continue
        m = re.match(r"Chapter\s+(\d+)", text_stripped)
        if m:
            ch_num = int(m.group(1))
            chapter_starts.append((ei, ch_num))

    for idx, (ei, ch_num) in enumerate(chapter_starts):
        if idx + 1 < len(chapter_starts):
            next_ei = chapter_starts[idx + 1][0]
            end_ei = next_ei
            for pi, pei, ptext in heading1_indices:
                if pei == next_ei - 1 and ptext.strip().startswith("PART"):
                    end_ei = pei
                    break
            sections[f"chapter-{ch_num}"] = elements[ei:end_ei]
        else:
            app_a_ei = None
            for eei, (etype, obj) in enumerate(elements):
                if etype == "paragraph" and obj.text.strip() == "APPENDIX A":
                    app_a_ei = eei
                    break
            if app_a_ei:
                sections[f"chapter-{ch_num}"] = elements[ei:app_a_ei]
            else:
                sections[f"chapter-{ch_num}"] = elements[ei:]

    app_a_ei = None
    app_b_ei = None
    for eei, (etype, obj) in enumerate(elements):
        if etype == "paragraph":
            if obj.text.strip() == "APPENDIX A":
                app_a_ei = eei
            elif obj.text.strip() == "APPENDIX B":
                app_b_ei = eei

    if app_a_ei and app_b_ei:
        sections["appendix-a"] = elements[app_a_ei:app_b_ei]
    elif app_a_ei:
        sections["appendix-a"] = elements[app_a_ei:]
    if app_b_ei:
        sections["appendix-b"] = elements[app_b_ei:]

    return sections


def para_has_content(paragraph):
    """Check if paragraph has any visible text or math."""
    if paragraph.text.strip():
        return True
    # Check for math elements
    return any(MATH_NS in c.tag for c in paragraph._element)


def section_to_html(section_elements, skip_count=0):
    """Convert a list of (type, obj) elements into HTML body content."""
    html_parts = []
    skipped = 0
    has_math = False

    for etype, obj in section_elements:
        if etype == "table":
            html_parts.append(table_to_html(obj))
            continue

        p = obj
        if not para_has_content(p):
            continue

        if skipped < skip_count:
            skipped += 1
            continue

        # Heading 1
        if p.style.name.startswith("Heading"):
            html_parts.append(f"<h1>{runs_to_html(p)}</h1>")
            continue

        # Numbered sub-section headings
        text = p.text.strip()
        if is_section_heading(p):
            level = classify_heading_level(text)
            content, _ = paragraph_to_html(p)
            html_parts.append(f"<{level}>{content}</{level}>")
            continue

        # Regular paragraph (with potential math)
        content, is_display_math = paragraph_to_html(p)

        if is_display_math:
            # Display math — wrap in a div, not a <p>
            html_parts.append(f'<div class="math-display">{content}</div>')
            has_math = True
        else:
            if "\\(" in content:
                has_math = True
            if text.startswith("\u2022") or text.startswith("- "):
                html_parts.append(f'<p class="list-item">{content}</p>')
            else:
                html_parts.append(f"<p>{content}</p>")

    return "\n".join(html_parts), has_math


# ── Chapter metadata ────────────────────────────────────────────────

CHAPTER_TITLES = {
    0: "A Universe of Motion", 1: "Heat", 2: "Polarity", 3: "Existence",
    4: "Righteousness", 5: "Order", 6: "Movement", 7: "Entropy",
    8: "Learning Systems", 9: "Identity and Persistence",
    10: "Agency and Choice", 11: "Error, Correction, and Growth",
    12: "Ethics as Stability", 13: "Coercion as Forced Motion",
    14: "Freedom as Available Motion",
}

PARTS = {
    0: ("I", "Orientation"),
    1: ("II", "The Motion Calendar"), 2: ("II", "The Motion Calendar"),
    3: ("II", "The Motion Calendar"), 4: ("II", "The Motion Calendar"),
    5: ("II", "The Motion Calendar"), 6: ("II", "The Motion Calendar"),
    7: ("III", "Systems"), 8: ("III", "Systems"), 9: ("III", "Systems"),
    10: ("III", "Systems"), 11: ("III", "Systems"),
    12: ("IV", "Meaning"), 13: ("IV", "Meaning"), 14: ("IV", "Meaning"),
}

NAV_ORDER = (["preface"] + [f"chapter-{i}" for i in range(15)]
             + ["appendix-a", "appendix-b"])

NAV_LABELS = {"preface": "Preface", "appendix-a": "Appendix A",
              "appendix-b": "Appendix B"}
for i in range(15):
    NAV_LABELS[f"chapter-{i}"] = f"Chapter {i}"


# ── HTML Templates ──────────────────────────────────────────────────

KATEX_HEAD = f"""<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.css">
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/katex.min.js"></script>
<script defer src="https://cdn.jsdelivr.net/npm/katex@0.16.11/dist/contrib/auto-render.min.js"
  onload="renderMathInElement(document.body, {{
    delimiters: [
      {{left: '\\\\[', right: '\\\\]', display: true}},
      {{left: '\\\\(', right: '\\\\)', display: false}}
    ],
    throwOnError: false
  }});"></script>"""


def page_html(title, body_content, page_id, subtitle=None, needs_math=False):
    """Wrap body content in a full HTML page."""
    idx = NAV_ORDER.index(page_id) if page_id in NAV_ORDER else -1
    prev_link = ""
    next_link = ""
    if idx > 0:
        prev_id = NAV_ORDER[idx - 1]
        prev_link = f'<a href="{BASE_URL}/{prev_id}.html">&larr; {NAV_LABELS[prev_id]}</a>'
    if 0 <= idx < len(NAV_ORDER) - 1:
        next_id = NAV_ORDER[idx + 1]
        next_link = f'<a href="{BASE_URL}/{next_id}.html">{NAV_LABELS[next_id]} &rarr;</a>'

    subtitle_html = f'<p class="subtitle">{escape(subtitle)}</p>' if subtitle else ""
    math_head = KATEX_HEAD if needs_math else ""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{escape(title)} — Ocean From Motion</title>
<link rel="stylesheet" href="{BASE_URL}/style.css">
{math_head}
</head>
<body>
<header>
  <nav class="top-nav">
    <a href="{BASE_URL}/" class="nav-title">Ocean From Motion</a>
    <div class="nav-links">
      <a href="{BASE_URL}/">Contents</a>
      <a href="{BASE_URL}/Ocean_From_Motion.docx" class="download-link">Download .docx</a>
    </div>
  </nav>
</header>
<main>
  <h1 class="page-title">{escape(title)}</h1>
  {subtitle_html}
  {body_content}
</main>
<footer>
  <nav class="chapter-nav">
    <div class="nav-prev">{prev_link}</div>
    <div class="nav-toc"><a href="{BASE_URL}/">Table of Contents</a></div>
    <div class="nav-next">{next_link}</div>
  </nav>
</footer>
</body>
</html>
"""


def index_html():
    """Generate the cover/index page."""
    toc_items = []
    current_part = None
    for ch_num in range(15):
        part_num, part_name = PARTS[ch_num]
        part_key = f"PART {part_num}"
        if part_key != current_part:
            current_part = part_key
            toc_items.append(f'<li class="part-header">Part {part_num}: {part_name}</li>')
        toc_items.append(
            f'<li><a href="{BASE_URL}/chapter-{ch_num}.html">'
            f"Chapter {ch_num}: {CHAPTER_TITLES[ch_num]}</a></li>"
        )
    toc_items.append('<li class="part-header">Appendices</li>')
    toc_items.append(f'<li><a href="{BASE_URL}/appendix-a.html">Appendix A: Mathematical Constants</a></li>')
    toc_items.append(f'<li><a href="{BASE_URL}/appendix-b.html">Appendix B: Notation Reference</a></li>')

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Ocean From Motion</title>
<link rel="stylesheet" href="{BASE_URL}/style.css">
</head>
<body>
<main class="cover">
  <h1 class="book-title">Ocean From Motion</h1>
  <p class="book-subtitle">A Study in the Nature of Possible Primitives of Motion</p>
  <p class="book-subtitle-secondary">Incorporating<br>The Motion Calendar<br>
  <em>A Universe of Motion Rather Than in Motion</em></p>
  <div class="download-section">
    <a href="{BASE_URL}/Ocean_From_Motion.docx" class="btn-download">Download Original (.docx)</a>
  </div>
  <nav class="toc">
    <h2>Contents</h2>
    <ol>
      <li><a href="{BASE_URL}/preface.html">Preface</a></li>
      {"".join(toc_items)}
    </ol>
  </nav>
</main>
</body>
</html>
"""


# ── Main ─────────────────────────────────────────────────────────────

def main():
    os.makedirs(DOCS_DIR, exist_ok=True)

    print("Reading .docx ...")
    doc = Document(DOCX_PATH)
    elements = parse_document(doc)
    sections = split_into_sections(elements)

    # Index
    print("Writing index.html")
    with open(os.path.join(DOCS_DIR, "index.html"), "w", encoding="utf-8") as f:
        f.write(index_html())

    # Preface
    print("Writing preface.html")
    preface_body, preface_math = section_to_html(sections["preface"], skip_count=1)
    with open(os.path.join(DOCS_DIR, "preface.html"), "w", encoding="utf-8") as f:
        f.write(page_html("Preface", preface_body, "preface", needs_math=preface_math))

    # Chapters
    for ch_num in range(15):
        key = f"chapter-{ch_num}"
        print(f"Writing {key}.html")
        if key not in sections:
            print(f"  WARNING: {key} not found!")
            continue

        ch_elements = sections[key]
        subtitle = None
        for etype, obj in ch_elements:
            if etype == "paragraph" and not obj.style.name.startswith("Heading"):
                if para_has_content(obj):
                    subtitle = obj.text.strip()
                    break

        body, has_math = section_to_html(ch_elements, skip_count=2)
        title = f"Chapter {ch_num}: {CHAPTER_TITLES[ch_num]}"

        with open(os.path.join(DOCS_DIR, f"{key}.html"), "w", encoding="utf-8") as f:
            f.write(page_html(title, body, key, subtitle=subtitle, needs_math=has_math))

    # Appendices
    for app_key, app_title in [("appendix-a", "Appendix A: Mathematical Constants"),
                                ("appendix-b", "Appendix B: Notation Reference")]:
        print(f"Writing {app_key}.html")
        if app_key not in sections:
            print(f"  WARNING: {app_key} not found!")
            continue
        app_body, app_math = section_to_html(sections[app_key], skip_count=2)
        with open(os.path.join(DOCS_DIR, f"{app_key}.html"), "w", encoding="utf-8") as f:
            f.write(page_html(app_title, app_body, app_key, needs_math=app_math))

    # Copy .docx
    dest = os.path.join(DOCS_DIR, "Ocean_From_Motion.docx")
    shutil.copy2(DOCX_PATH, dest)
    print(f"Copied .docx to {dest}")
    print("Done!")


if __name__ == "__main__":
    main()
